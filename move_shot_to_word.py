from __future__ import annotations

import os
import re
import base64
import mimetypes
from html import escape
from pathlib import Path
from typing import Dict, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm
from dotenv import load_dotenv

from PIL import Image  # NEW: для изменения размеров изображений

# COM-экспорт Word -> HTML (опционально, только Windows + установлен Word)
try:
    import win32com.client as win32  # type: ignore

    _HAS_WORD = True
except Exception:
    _HAS_WORD = False

from config_app import (
    SCREENSHOTS_DIR,
    CURRENT_MONTH,
    CURRENT_YEAR,
    BASE_URL_TH,
    BASE_URL_PRO,
)
from utils import (
    get_desktop_dir,
    normalize_windows_path,
    load_links,
)  # структура/пути как у вас

# =========================
# Константы и настройки
# =========================

FONT_NAME = "Arial"
FONT_SIZE_TITLE = 14
FONT_SIZE_HOTEL_LINK = 18
FONT_SIZE_CAPTION = 12

IMAGE_WIDTH_INCHES = 6  # ширина вставки в DOCX (это не пиксели файла; файл мы теперь можем заранее привести)
PAGE_BREAK_FILES = {"07_rating_in_hurghada.png", "08_activity.png", "04_attendance.png"}

URL_RE = re.compile(r"(https?://[^\s)]+)")

load_dotenv()

# =========================
# Помощники для DOCX (как у вас, без изменений в семантике)
# =========================


def add_header_image(
    doc: Document,
    image_path: str | Path,
    width_inches: float = 2.0,
    top_margin_cm: float = 1.0,
    bottom_margin_cm: float = 0.5,
) -> None:
    image_path = Path(image_path)
    if not image_path.exists():
        raise FileNotFoundError(f"Файл {image_path} не найден.")
    for section in doc.sections:
        section.header_distance = Cm(top_margin_cm)
        header = section.header
        paragraph = (
            header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))
        if bottom_margin_cm > 0:
            empty_p = header.add_paragraph()
            empty_p.paragraph_format.space_before = Pt(0)
            empty_p.paragraph_format.space_after = Cm(bottom_margin_cm)


def ensure_normal_style_arial(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    rPr = style._element.rPr
    rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def set_run_arial(run, size_pt: int) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    rPr = run._element.rPr
    rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def add_hyperlink(
    paragraph,
    text: str,
    url: str,
    *,
    font_name: str = FONT_NAME,
    font_size_pt: int = FONT_SIZE_CAPTION,
    bold: bool = False,
    underline: bool = True,
    color: str = "0000FF",
):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.font.bold = bold
    run.font.underline = underline
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.rPr
    rPr.rFonts.set(qn("w:ascii"), font_name)
    rPr.rFonts.set(qn("w:hAnsi"), font_name)
    rPr.rFonts.set(qn("w:eastAsia"), font_name)

    hyperlink.append(run._element)
    paragraph._p.append(hyperlink)
    return run


def add_text_with_links(
    paragraph, text: str, *, font_size_pt: int = FONT_SIZE_CAPTION
) -> None:
    pos = 0
    for m in URL_RE.finditer(text or ""):
        if m.start() > pos:
            prefix = text[pos : m.start()]
            if prefix:
                set_run_arial(paragraph.add_run(prefix), size_pt=font_size_pt)
        url = m.group(1)
        add_hyperlink(paragraph, url, url, font_size_pt=font_size_pt)
        pos = m.end()
    if pos < len(text or ""):
        tail = text[pos:]
        if tail:
            set_run_arial(paragraph.add_run(tail), size_pt=font_size_pt)


def build_mapping(
    hotel_id: int, *, rating_url: str | None = None, city: str, star: str
) -> Dict[str, str]:
    base_url_pro_without_ssa = BASE_URL_PRO.replace("ssa.", "")
    base = base_url_pro_without_ssa + "hotel/" + str(hotel_id)
    rating_url = rating_url or f"{base}/new_stat/rating-hotels"
    return {
        "01_top_element.png": "",
        "02_populars_element.png": "Popularity of the hotel",
        "03_reviews.png": "Rating and recommendations of hotel",
        "04_attendance.png": f"Hotel profile attendance by month: {base}/new_stat/attendance",
        "06_service_prices.png": f"Log of booking requests: {base}/stat/profile?group=week&vw=grouped",
        "07_rating_in_hurghada.png": f"Ranking of {city} {star} hotels by ratings over the last 2 years: {rating_url}",
        "08_activity.png": f"Last page activity: {base}/activity/index",
    }


def build_reports_dir(curr_year: str, curr_month: str, city: str) -> Path:
    raw_path = os.getenv("PATH_FOR_REPORTS")
    base_dir = normalize_windows_path(raw_path) if raw_path else get_desktop_dir()
    reports = (
        Path(base_dir)
        / "TopHotels Reports"
        / f"{curr_year}"
        / f"{curr_month}"
        / f"{city}"
    )
    reports.mkdir(parents=True, exist_ok=True)
    return reports


# =========================
# HTML (резервный чистый inline, как было у вас)
# =========================


def _linkify(text: str) -> str:
    if not text:
        return ""
    return re.sub(
        r"(https?://[^\s)]+)", r'<a href="\1" target="_blank">\1</a>', escape(text)
    )


def _img_to_data_uri(path: Path) -> str:
    mime, _ = mimetypes.guess_type(path.name)
    if not mime:
        mime = "application/octet-stream"
    data = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:{mime};base64,{data}"


def _build_inline_html(title_hotel: str, url_hotel: str,
                       mapping_paragraph: Dict[str, str], folder_path: Path) -> str:
    max_width_px = int(IMAGE_WIDTH_INCHES * 96)
    css = f"""
    body{{font-family:Arial,Helvetica,sans-serif; color:#111;}}
    .wrap{{max-width:{max_width_px + 40}px; margin:24px auto;}}
    h1, h2, h3{{margin:8px 0; text-align:center}}
    a{{color:#0645AD;}}
    /* список подписей */
    ul.caplist{{
        margin-top:20px;
        margin-bottom:20px;
        padding-left:26px;
    }}
    ul.caplist > li{{
        margin-top:8px;
        margin-bottom:8px;
        font-size:12pt;
        line-height:1.5;
    }}
    /* картинка */
    .imgbox{{
        margin-top:20px;
        margin-bottom:30px;
    }}
    img{{
        max-width:100%;
        height:auto;
        display:block;
        border:1px solid #ddd;
    }}
    """

    parts = []
    parts.append(
        f"<!doctype html><html><head><meta charset='utf-8'>"
        f"<meta name='viewport' content='width=device-width,initial-scale=1'>"
        f"<style>{css}</style></head><body><div class='wrap'>"
    )
    parts.append(f"<h2>Monthly Statistics Report {escape(CURRENT_MONTH)} {escape(CURRENT_YEAR)}</h2>")
    parts.append(f"<h1><a href='{escape(url_hotel)}' target='_blank'>{escape(title_hotel.upper())}</a></h1>")

    for file_name in sorted(os.listdir(folder_path)):
        if not file_name.lower().endswith((".png", ".jpg", ".jpeg")):
            continue

        # --- подпись как UL/LI с отступами сверху/снизу ---
        caption = mapping_paragraph.get(file_name)
        if caption is not None:
            if caption:
                parts.append(
                    f"<ul style='margin:20px 0; padding-left:26px; font-size:12pt; line-height:1.5; font-family:Arial,Helvetica,sans-serif;'>"
                )
                parts.append(
                    f"<li style='margin:8px 0; font-size:12pt; line-height:1.5; font-family:Arial,Helvetica,sans-serif;'>{_linkify(caption)}</li>"
                )
                parts.append("</ul>")

        # --- картинка ---
        img_path = folder_path / file_name
        try:
            data_uri = _img_to_data_uri(img_path)
            parts.append(f"<div class='imgbox'><img src='{data_uri}' alt=''></div>")
        except Exception as e:
            parts.append(f"<div style='color:#b00; font-size:12pt'>[Image error: {escape(str(e))}]</div>")

    parts.append("</div></body></html>")
    return "".join(parts)



# =========================
# NEW: Приведение картинок к фиксированной ширине (px)
# =========================


def _resize_all_images(folder_path: Path, width_px: int) -> None:
    """
    Пройтись по всем .png/.jpg/.jpeg в папке и привести их к фиксированной ширине (px),
    сохраняя пропорции. Перезапись *на месте* (in-place).
    """
    if width_px <= 0:
        return
    for name in sorted(os.listdir(folder_path)):
        if not name.lower().endswith((".png", ".jpg", ".jpeg")):
            continue
        p = folder_path / name
        try:
            with Image.open(p) as im:
                w, h = im.size
                if w == width_px:
                    continue
                scale = width_px / float(w)
                new_h = max(1, int(round(h * scale)))
                im = im.resize((width_px, new_h), Image.LANCZOS)
                # сохраняем с тем же форматом, максимально без потери видимого качества
                if p.suffix.lower() in (".jpg", ".jpeg"):
                    im.save(p, quality=95, optimize=True, progressive=True)
                else:
                    # для PNG — без потерь, с оптимизацией
                    im.save(p, optimize=True)
        except Exception as e:
            print(f"[WARN] Resize failed for {p.name}: {e}")


# =========================
# NEW: Экспорт DOCX -> Word HTML и встраивание картинок
# =========================


def _export_docx_to_word_html_and_inline(docx_path: Path) -> Optional[Path]:
    """
    Экспортирует DOCX в 'Filtered HTML' через COM (Word), затем:
    - читает созданный HTML
    - находит <img src="..."> и подменяет src на data:URI (base64), читая файлы из ..._files/
    - сохраняет рядом *_word_inline.html
    Возвращает путь к inline-HTML или None, если COM недоступен или произошла ошибка.
    """
    if not _HAS_WORD:
        return None
    try:
        docx_path = Path(docx_path)
        html_path = docx_path.with_suffix(".htm")  # Word по умолчанию .htm
        assets_dir = docx_path.with_name(docx_path.stem + "_files")  # папка ресурсов

        word = win32.gencache.EnsureDispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path))
        # 10 = wdFormatFilteredHTML (более «чистый» HTML с mso-стилями, пригодный для Outlook)
        doc.SaveAs(str(html_path), FileFormat=10)
        doc.Close(False)
        word.Quit()

        # Заменим <img src="assets/..."> на data: URI
        html = html_path.read_text(encoding="utf-8", errors="ignore")

        def _to_data_uri(img_rel_src: str) -> str:
            img_path = assets_dir / Path(img_rel_src).name
            if not img_path.exists():
                # иногда Word пишет относительные пути с подпапками — попробуем как есть
                img_path = assets_dir / img_rel_src
                if not img_path.exists():
                    return img_rel_src  # оставим как есть
            mime, _ = mimetypes.guess_type(img_path.name)
            if not mime:
                mime = "application/octet-stream"
            b64 = base64.b64encode(img_path.read_bytes()).decode("ascii")
            return f"data:{mime};base64,{b64}"

        # грубая, но практичная замена src в тегах <img ...>
        html_inline = re.sub(
            r'(<img\b[^>]*\bsrc=")([^"]+)(")',
            lambda m: f"{m.group(1)}{_to_data_uri(m.group(2))}{m.group(3)}",
            html,
            flags=re.IGNORECASE,
        )

        out_inline = docx_path.with_name(docx_path.stem + "_word_inline.html")
        out_inline.write_text(html_inline, encoding="utf-8")

        # (опционально) можно удалить html и папку ресурсов; я оставлю их, чтобы можно было сравнить
        print(f"✔ Word Filtered HTML (inline) created: {out_inline}")
        return out_inline
    except Exception as e:
        print(f"[WARN] Word HTML export failed: {e}")
        return None


# =========================
# Основная логика
# =========================


def create_formatted_doc(target_image_width_px: int | None = None) -> None:
    """
    Генерирует DOCX и HTML-версии отчёта.
    Дополнительно: если задан target_image_width_px, сначала приводит каждую картинку
    в папке отеля к фиксированной ширине (px) с сохранением пропорций (in-place).
    """
    screenshots_dir = Path(SCREENSHOTS_DIR)

    for folder_name in os.listdir(screenshots_dir):
        if "_" not in folder_name:
            continue  # пропускаем «посторонние» папки
        hotel_id, title_hotel = folder_name.split("_", 1)
        folder_path = screenshots_dir / folder_name
        if not folder_path.is_dir():
            continue

        # 1) Привести картинки к фиксированной ширине, если задано
        if isinstance(target_image_width_px, int) and target_image_width_px > 0:
            _resize_all_images(folder_path, target_image_width_px)

        # 2) Метаданные и директории отчётов
        json_file = load_links(hotel_id, title_hotel)
        city = json_file.get("city", "City")
        reports_dir = build_reports_dir(CURRENT_YEAR, CURRENT_MONTH, city)
        star = json_file.get("star", "*")
        rating_url = json_file.get("rating_url")
        url_hotel = f"{BASE_URL_TH}hotel/{hotel_id}"
        mapping_paragraph = build_mapping(
            hotel_id, rating_url=rating_url, city=city, star=star
        )

        # 3) DOCX
        doc = Document()
        add_header_image(
            doc, "th_logo/logo_1.jpg", width_inches=1.5, bottom_margin_cm=0.2
        )
        ensure_normal_style_arial(doc)

        title_1 = doc.add_paragraph()
        title_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_1 = title_1.add_run(
            f"Monthly Statistics Report {CURRENT_MONTH} {CURRENT_YEAR}"
        )
        set_run_arial(run_1, size_pt=FONT_SIZE_TITLE)

        title_2 = doc.add_paragraph()
        title_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_hyperlink(
            title_2,
            title_hotel.upper(),
            url_hotel,
            font_size_pt=FONT_SIZE_HOTEL_LINK,
            bold=True,
        )

        for file_name in sorted(os.listdir(folder_path)):
            if not file_name.lower().endswith((".png", ".jpg", ".jpeg")):
                continue
            if file_name in PAGE_BREAK_FILES:
                doc.add_page_break()

            caption = mapping_paragraph.get(file_name)
            if caption is not None:
                p = doc.add_paragraph(style="List Bullet")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # добавляем текст
                if caption:
                    add_text_with_links(p, caption, font_size_pt=FONT_SIZE_CAPTION)

                # задаём отступы сверху и снизу
                p.paragraph_format.space_before = Pt(6)  # пробел сверху
                p.paragraph_format.space_after = Pt(6)  # пробел снизу

            image_path = folder_path / file_name
            try:
                # Вставка с фиксированной шириной страницы (как и было)
                doc.add_picture(str(image_path), width=Inches(IMAGE_WIDTH_INCHES))
            except Exception as e:
                err_p = doc.add_paragraph(f"[Image error: {e}]")
                for r in err_p.runs:
                    set_run_arial(r, size_pt=FONT_SIZE_CAPTION)

        safe_name = title_hotel.replace(" ", "_").replace("*", "")
        docx_path = reports_dir / f"{safe_name}.docx"
        doc.save(docx_path)
        print(f"✔ Report created: {docx_path}")

        # 4) HTML: сначала пробуем Word Filtered HTML с интегрированными картинками (сохранение верстки Word)
        word_inline_html = _export_docx_to_word_html_and_inline(docx_path)

        if word_inline_html is None:
            # 5) Fallback: ваш старый «чистый» inline-HTML (без вордовских mso-стилей)
            html = _build_inline_html(
                title_hotel, url_hotel, mapping_paragraph, folder_path
            )
            html_path = reports_dir / f"{safe_name}_inline.html"
            html_path.write_text(html, encoding="utf-8")
            print(f"✔ Inline HTML (fallback) created: {html_path}")


if __name__ == "__main__":
    # Пример: привести все скриншоты к ширине 1200px перед сборкой документов
    create_formatted_doc(target_image_width_px=1200)
