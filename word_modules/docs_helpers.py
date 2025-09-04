from __future__ import annotations

import os
from pathlib import Path
from typing import Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm

from config_app import BASE_URL_PRO, FONT_SIZE_CAPTION, FONT_NAME, URL_RE
from utils import normalize_windows_path, get_desktop_dir


def add_header_image(
    doc: Document,
    image_path: str | Path,
    width_inches: float = 2.0,
    top_margin_cm: float = 1.0,
    bottom_margin_cm: float = 0.5,
) -> None:
    image_path = Path(image_path)
    if not image_path.exists():
        raise FileNotFoundError(f"Файл {image_path} не найден.")
    for section in doc.sections:
        section.header_distance = Cm(top_margin_cm)
        header = section.header
        paragraph = (
            header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))
        if bottom_margin_cm > 0:
            empty_p = header.add_paragraph()
            empty_p.paragraph_format.space_before = Pt(0)
            empty_p.paragraph_format.space_after = Cm(bottom_margin_cm)


def ensure_normal_style_arial(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    rPr = style._element.rPr
    rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def set_run_arial(run, size_pt: int) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    rPr = run._element.rPr
    rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def add_hyperlink(
    paragraph,
    text: str,
    url: str,
    *,
    font_name: str = FONT_NAME,
    font_size_pt: int = FONT_SIZE_CAPTION,
    bold: bool = False,
    underline: bool = True,
    color: str = "0000FF",
):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.font.bold = bold
    run.font.underline = underline
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.rPr
    rPr.rFonts.set(qn("w:ascii"), font_name)
    rPr.rFonts.set(qn("w:hAnsi"), font_name)
    rPr.rFonts.set(qn("w:eastAsia"), font_name)

    hyperlink.append(run._element)
    paragraph._p.append(hyperlink)
    return run


def add_text_with_links(
    paragraph, text: str, *, font_size_pt: int = FONT_SIZE_CAPTION
) -> None:
    pos = 0
    for m in URL_RE.finditer(text or ""):
        if m.start() > pos:
            prefix = text[pos : m.start()]
            if prefix:
                set_run_arial(paragraph.add_run(prefix), size_pt=font_size_pt)
        url = m.group(1)
        add_hyperlink(paragraph, url, url, font_size_pt=font_size_pt)
        pos = m.end()
    if pos < len(text or ""):
        tail = text[pos:]
        if tail:
            set_run_arial(paragraph.add_run(tail), size_pt=font_size_pt)


def build_mapping(
    hotel_id: int, *, rating_url: str | None = None, city: str, star: str
) -> Dict[str, str]:
    base_url_pro_without_ssa = BASE_URL_PRO.replace("ssa.", "")
    base = base_url_pro_without_ssa + "hotel/" + str(hotel_id)
    rating_url = rating_url or f"{base}/new_stat/rating-hotels"
    return {
        "01_top_element.png": "",
        "02_populars_element.png": "Popularity of the hotel\n",
        "03_reviews.png": f"Rating and recommendations \n",
        "04_attendance.png": f"Hotel profile attendance by month: {base}/new_stat/attendance\n",
        "06_service_prices.png": f"Log of booking requests: {base}/stat/profile?group=week&vw=grouped\n",
        "07_rating_in_hurghada.png": f"Ranking of {city} {star} for the last 2 years: {rating_url}\n",
        "08_activity.png": f"Last month's activities: {base}/activity/index\n",
    }


def build_reports_dir(curr_year: str, curr_month: str, city: str, chain:str) -> Path:
    raw_path = os.getenv("PATH_FOR_REPORTS")
    base_dir = normalize_windows_path(raw_path) if raw_path else get_desktop_dir()
    reports = (
        Path(base_dir)
        / "TopHotels Reports"
        / f"{curr_year}"
        / f"{curr_month}"
        / f"{city}"
        / f"{chain}"
    )
    reports.mkdir(parents=True, exist_ok=True)
    return reports
