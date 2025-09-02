from __future__ import annotations

import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.shared import Inches, Pt

from wor_modules.docs_helpers import (
    add_header_image,
    ensure_normal_style_arial,
    set_run_arial,
    add_hyperlink,
    add_text_with_links,
)

try:
    import win32com.client as win32  # type: ignore

    _HAS_WORD = True
except Exception:
    _HAS_WORD = False

from config_app import (
    CURRENT_MONTH,
    CURRENT_YEAR,
    FONT_SIZE_TITLE,
    FONT_SIZE_HOTEL_LINK,
    FONT_SIZE_CAPTION,
    IMAGE_WIDTH_INCHES,
    PAGE_BREAK_FILES,
)


def create_word_file(
    title_hotel, folder_path, url_hotel, mapping_paragraph, reports_dir
):
    # 3) DOCX
    doc = Document()
    add_header_image(doc, "th_logo/logo_1.jpg", width_inches=1.5, bottom_margin_cm=0.2)
    ensure_normal_style_arial(doc)

    title_1 = doc.add_paragraph()
    title_1.add_run("Hi, name.\n")
    run_1 = title_1.add_run(
        f"Monthly Statistics Report {CURRENT_MONTH} {CURRENT_YEAR}  "
    )
    set_run_arial(run_1, size_pt=FONT_SIZE_TITLE)

    add_hyperlink(
        title_1,
        title_hotel.upper(),
        url_hotel,
        font_size_pt=FONT_SIZE_HOTEL_LINK,
        bold=True,
    )

    for file_name in sorted(os.listdir(folder_path)):
        if not file_name.lower().endswith((".png", ".jpg", ".jpeg")):
            continue
        if file_name in PAGE_BREAK_FILES:
            doc.add_page_break()

        caption = mapping_paragraph.get(file_name)
        if caption is not None:
            p = doc.add_paragraph(style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # добавляем текст
            if caption:
                add_text_with_links(p, caption, font_size_pt=FONT_SIZE_CAPTION)

            # задаём отступы сверху и снизу
            p.paragraph_format.space_before = Pt(6)  # пробел сверху
            p.paragraph_format.space_after = Pt(6)  # пробел снизу

        image_path = folder_path / file_name
        try:
            # Вставка с фиксированной шириной страницы (как и было)
            doc.add_picture(str(image_path), width=Inches(IMAGE_WIDTH_INCHES))
        except Exception as e:
            err_p = doc.add_paragraph(f"[Image error: {e}]")
            for r in err_p.runs:
                set_run_arial(r, size_pt=FONT_SIZE_CAPTION)

    safe_name = title_hotel.replace(" ", "_").replace("*", "")
    docx_path = reports_dir / f"{safe_name}.docx"
    doc.save(docx_path)
    print(f"✔ Report created: {docx_path}")
